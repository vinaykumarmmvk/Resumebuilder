from docx import Document
from docx.shared import Inches
from docx.shared import Cm

pic = "../dp_vinay.jpeg"
document = Document()
tbl = document.add_table(rows=1, cols=2)
tbl.autofit = False 
tbl.allow_autofit = False
tbl.cell(0,0).width = Inches(5.0) 
tbl.cell(0,1).width = Inches(1.5) 

row_cells = tbl.add_row().cells
row_cells[0].width = Inches(1.0)
row_cells[1].width = Inches(1.0)
paragraph = row_cells[0].paragraphs[0]
run = paragraph.add_run()
run.add_text("abcdef")

paragraph1 = row_cells[1].paragraphs[0]
run1 = paragraph1.add_run()
#run1.add_picture(pic, width = 1000000, height = 1000000)


document.save("demo.docx")