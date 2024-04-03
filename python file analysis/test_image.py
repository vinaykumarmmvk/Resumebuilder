from docx import Document
from docx.shared import Inches

doc = Document('addImage1.docx')
#tables = doc.tables
tables = doc.add_table(rows=1, cols=3)
p = tables[0].rows[0].cells[0].add_paragraph()
r = p.add_run()
r.add_picture('dp_vinay.jpeg',width=Inches(4.0), height=Inches(.7))
p = tables[1].rows[0].cells[0].add_paragraph()
r = p.add_run()
r.add_picture('dp_vinay.jpeg',width=Inches(4.0), height=Inches(.7))
doc.save('addImage1.docx')

